import json
import os
import logging
import requests
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from typing import Dict, List, Optional

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
HUGGINGFACE_API_KEY = os.getenv('HUGGINGFACE_API_KEY')
if not HUGGINGFACE_API_KEY:
    logger.error("HUGGINGFACE_API_KEY not found in environment variables")
    raise ValueError("HUGGINGFACE_API_KEY is required")

# Load master resume
def load_master_resume() -> Dict:
    """Load the master resume from JSON file."""
    try:
        with open('master_resume.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        logger.error("master_resume.json not found")
        raise
    except json.JSONDecodeError as e:
        logger.error(f"Error parsing master_resume.json: {e}")
        raise

def generate_resume_content(master_resume: Dict, job_details: Dict) -> str:
    """Generate ATS-optimized resume content using Hugging Face Llama 3.1 8B-Instruct."""
    try:
        prompt = f"""
        Generate an ATS-optimized resume for the following job:

        Job Title: {job_details.get('job_title', 'Unknown')}
        Required Skills: {', '.join(job_details.get('skills', []))}
        Employer: {job_details.get('employer_email', 'Unknown')}

        Base Resume Information:
        {json.dumps(master_resume, indent=2)}

        Instructions:
        1. Tailor the resume to match the job requirements
        2. Include relevant keywords from the job skills
        3. Keep the format clean and ATS-friendly
        4. Do not fabricate any experience or skills
        5. Focus on quantifiable achievements
        6. Keep it concise (1 page)
        7. Format as:
           - Header: Name, Email, Phone, Location (plain text)
           - Summary: 2-3 sentences with job keywords
           - Skills: Bulleted list
           - Experience: Reverse chronological, job-relevant
           - Education: Degree and institution
           - Certifications: Job-relevant ones

        Return only the resume content in plain text format.
        """

        # Hugging Face Inference API call
        response = requests.post(
            'https://api-inference.huggingface.co/models/meta-llama/Llama-3.1-8B-Instruct',
            headers={
                'Authorization': f'Bearer {HUGGINGFACE_API_KEY}',
                'Content-Type': 'application/json'
            },
            json={
                'inputs': prompt,
                'parameters': {
                    'max_new_tokens': 2000,
                    'temperature': 0.7,
                    'do_sample': True
                }
            },
            timeout=60  # Increased timeout for inference
        )

        if response.status_code == 200:
            result = response.json()
            if isinstance(result, list) and result:
                return result[0].get('generated_text', '').strip()
            else:
                raise Exception("Unexpected API response format")
        else:
            logger.error(f"Hugging Face API error: {response.status_code} - {response.text}")
            raise Exception(f"Failed to generate resume: {response.status_code}")

    except requests.RequestException as e:
        logger.error(f"Request error: {e}")
        raise
    except Exception as e:
        logger.error(f"Error generating resume content: {e}")
        raise

def create_word_resume(content: str, filename: str) -> str:
    """Create an ATS-friendly Word document from the resume content."""
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document()

        # Set default font to Arial 11pt
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        lines = content.split('\n')
        current_section = ""

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect sections
            if line.upper() in ['SUMMARY', 'SKILLS', 'EXPERIENCE', 'EDUCATION', 'CERTIFICATIONS']:
                current_section = line.upper()
                p = doc.add_paragraph()
                p.add_run(line).bold = True
                continue

            p = doc.add_paragraph()

            # Handle bullets for skills and certifications
            if current_section in ['SKILLS', 'CERTIFICATIONS'] and (line.startswith('-') or line.startswith('•')):
                p.style = 'List Bullet'
                p.add_run(line[1:].strip())
            else:
                p.add_run(line)

        filepath = f"{filename}.docx"
        doc.save(filepath)
        logger.info(f"ATS-friendly Word resume saved to {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        raise

def create_pdf_resume(content: str, filename: str) -> str:
    """Create an ATS-friendly PDF document from the resume content."""
    try:
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_LEFT

        filepath = f"{filename}.pdf"
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()

        # Create ATS-friendly style
        ats_style = ParagraphStyle(
            'ATS',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            alignment=TA_LEFT
        )

        story = []
        lines = content.split('\n')
        current_section = ""

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect sections
            if line.upper() in ['SUMMARY', 'SKILLS', 'EXPERIENCE', 'EDUCATION', 'CERTIFICATIONS']:
                current_section = line.upper()
                story.append(Paragraph(line, ats_style))
                story.append(Spacer(1, 6))
                continue

            # Handle bullets for skills and certifications
            if current_section in ['SKILLS', 'CERTIFICATIONS'] and (line.startswith('-') or line.startswith('•')):
                bullet_style = ParagraphStyle(
                    'Bullet',
                    parent=ats_style,
                    leftIndent=20,
                    bulletIndent=10
                )
                story.append(Paragraph(f"• {line[1:].strip()}", bullet_style))
            else:
                story.append(Paragraph(line, ats_style))

            story.append(Spacer(1, 6))

        doc.build(story)
        logger.info(f"ATS-friendly PDF resume saved to {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Error creating PDF document: {e}")
        raise

def generate_resume(job_details: Dict) -> Dict:
    """Main function to generate ATS-optimized resume."""
    try:
        master_resume = load_master_resume()
        content = generate_resume_content(master_resume, job_details)

        # Generate unique filename based on job
        job_id = job_details.get('email_id', 'unknown')
        base_filename = f"resume_{job_id}"

        word_file = create_word_resume(content, base_filename)
        pdf_file = create_pdf_resume(content, base_filename)

        return {
            'content': content,
            'word_file': word_file,
            'pdf_file': pdf_file,
            'job_title': job_details.get('job_title', ''),
            'employer_email': job_details.get('employer_email', '')
        }
    except Exception as e:
        logger.error(f"Error generating resume: {e}")
        return {
            'error': str(e),
            'content': '',
            'word_file': '',
            'pdf_file': '',
            'job_title': job_details.get('job_title', ''),
            'employer_email': job_details.get('employer_email', '')
        }

def generate_resumes_for_jobs(parsed_jobs: List[Dict]) -> List[Dict]:
    """Generate resumes for multiple jobs."""
    resumes = []
    for job in parsed_jobs:
        resume = generate_resume(job)
        resumes.append(resume)
    return resumes